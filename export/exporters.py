"""
Data export module.

Supports exporting reports and events to:
  - Excel (.xlsx) via openpyxl
  - PDF via reportlab (text-based)
  - Word (.docx) via python-docx
  - Markdown (always supported)

Falls back gracefully if a library is not installed.
"""
from __future__ import annotations

import io
import json
from datetime import date, datetime
from pathlib import Path
from typing import Optional

from fastapi.responses import Response
from loguru import logger

from storage import get_db
from storage.models import DailyReport, IndustryEvent, KnowledgeSignal


def _has_openpyxl() -> bool:
    try:
        import openpyxl  # noqa
        return True
    except ImportError:
        return False


def _has_reportlab() -> bool:
    try:
        import reportlab  # noqa
        return True
    except ImportError:
        return False


def _has_docx() -> bool:
    try:
        from docx import Document  # noqa
        return True
    except ImportError:
        return False


# ============================================================================
# Excel export
# ============================================================================

def export_to_excel(sheets: dict[str, list[dict]], filename: str) -> Response:
    """Export sheets to Excel. `sheets` = {sheet_name: [rows]}."""
    if not _has_openpyxl():
        # Fallback to Markdown
        return _export_to_markdown(sheets, filename)

    from openpyxl import Workbook
    wb = Workbook()
    wb.remove(wb.active)
    for sheet_name, rows in sheets.items():
        ws = wb.create_sheet(title=sheet_name[:31])
        if not rows:
            ws.append(["(empty)"])
            continue
        # Header
        headers = list(rows[0].keys())
        ws.append(headers)
        # Rows
        for row in rows:
            ws.append([row.get(h, "") for h in headers])
        # Auto-size
        for col_idx, header in enumerate(headers, 1):
            col_letter = ws.cell(row=1, column=col_idx).column_letter
            ws.column_dimensions[col_letter].width = max(12, min(40, len(str(header)) + 2))

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}.xlsx"'},
    )


# ============================================================================
# PDF export (simple text-based)
# ============================================================================

def _register_chinese_font():
    """Register a CJK-capable font for reportlab. Returns the font name to use.

    Tries STSong-Light (built-in CID font, no external file needed) first.
    Falls back to Helvetica (Latin only) if registration fails.
    """
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light"
    except Exception as e:
        logger.warning(f"register STSong-Light failed: {e!r}")
    return "Helvetica"


_CJK_FONT = None


def _get_cjk_font() -> str:
    global _CJK_FONT
    if _CJK_FONT is None:
        _CJK_FONT = _register_chinese_font()
    return _CJK_FONT


def _markdown_inline_to_html(text: str) -> str:
    """Convert a single line of markdown inline syntax to reportlab-safe HTML.

    Escapes HTML-special chars first so reportlab doesn't choke on user content,
    then drops the markdown `**` and `` ` `` markers (kept simple to avoid
    breaking on edge cases).
    """
    from html import escape
    if not text:
        return ""
    s = escape(text)
    s = s.replace("**", "", 0)
    s = s.replace("`", "", 0)
    return s


def export_to_pdf(title: str, sections: list[tuple[str, list[dict]]], filename: str) -> Response:
    """Export to PDF using reportlab platypus. `sections` = [(heading, rows)].

    When a section's rows contain a single dict with a `content` key, the value
    is treated as full markdown body and rendered as flowing paragraphs
    (otherwise the whole body would get squashed into one unreadable table cell).
    """
    if not _has_reportlab():
        return _export_to_markdown(dict((h, r) for h, r in sections), filename)

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    )
    from reportlab.lib import colors

    cjk = _get_cjk_font()
    styles = getSampleStyleSheet()
    for sname in ("Title", "Heading1", "Heading2", "Heading3", "BodyText"):
        if sname in styles:
            styles[sname].fontName = cjk
    body_style = ParagraphStyle(
        "CJKBody", parent=styles["BodyText"], fontName=cjk, fontSize=9, leading=13,
    )
    h2_style = ParagraphStyle(
        "CJKH2", parent=styles["Heading2"], fontName=cjk, fontSize=12, leading=16,
    )
    title_style = ParagraphStyle(
        "CJKTitle", parent=styles["Title"], fontName=cjk, fontSize=16, leading=20,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, title=title, leftMargin=2*cm, rightMargin=2*cm)
    story = [Paragraph(_markdown_inline_to_html(title), title_style), Spacer(1, 0.4 * cm)]

    for heading, rows in sections:
        story.append(Paragraph(_markdown_inline_to_html(f"▎ {heading}"), h2_style))
        if not rows:
            story.append(Paragraph("<i>(no data)</i>", body_style))
            story.append(Spacer(1, 0.3 * cm))
            continue
        first = rows[0]
        if "content" in first and len(first) == 1:
            md = str(first.get("content") or "")
            for line in md.splitlines():
                stripped = line.strip()
                if not stripped:
                    story.append(Spacer(1, 0.15 * cm))
                    continue
                if stripped.startswith("### "):
                    story.append(Paragraph(_markdown_inline_to_html(stripped[4:]), h2_style))
                elif stripped.startswith("## "):
                    story.append(Paragraph(_markdown_inline_to_html(stripped[3:]), h2_style))
                elif stripped.startswith("# "):
                    story.append(Paragraph(_markdown_inline_to_html(stripped[2:]), title_style))
                elif stripped.startswith("|") and stripped.endswith("|"):
                    cells = [c.strip() for c in stripped.strip("|").split("|")]
                    story.append(Paragraph(" · ".join(_markdown_inline_to_html(c) for c in cells), body_style))
                elif stripped.startswith(("- ", "* ")):
                    story.append(Paragraph("• " + _markdown_inline_to_html(stripped[2:]), body_style))
                else:
                    story.append(Paragraph(_markdown_inline_to_html(stripped), body_style))
        else:
            headers = list(first.keys())
            data = [headers] + [
                [str(r.get(h, ""))[:80] for h in headers] for r in rows[:200]
            ]
            t = Table(data, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e40af")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, -1), cjk),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(t)
        story.append(Spacer(1, 0.3 * cm))

    doc.build(story)
    buf.seek(0)
    return Response(
        content=buf.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}.pdf"'},
    )


# ============================================================================
# Markdown export (always available fallback)
# ============================================================================

def _export_to_markdown(sheets: dict[str, list[dict]], filename: str) -> Response:
    """Render sheets as Markdown text."""
    lines = [f"# {filename}", ""]
    for sheet_name, rows in sheets.items():
        lines.append(f"## {sheet_name}")
        lines.append("")
        if not rows:
            lines.append("_(empty)_")
            lines.append("")
            continue
        headers = list(rows[0].keys())
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("|" + "|".join(["---"] * len(headers)) + "|")
        for row in rows:
            cells = [str(row.get(h, "")).replace("|", "\\|").replace("\n", " ")[:80] for h in headers]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")
    content = "\n".join(lines)
    return Response(
        content=content.encode("utf-8"),
        media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="{filename}.md"'},
    )


def export_to_markdown(sheets: dict[str, list[dict]], filename: str) -> Response:
    return _export_to_markdown(sheets, filename)


# ============================================================================
# High-level export functions
# ============================================================================

async def export_daily_report(
    report_date: str = None,
    format: str = "excel",
) -> Response:
    """Export a daily report to specified format."""
    db = get_db()
    target_date = None
    if report_date:
        try:
            target_date = date.fromisoformat(report_date)
        except ValueError:
            pass

    with db.session() as s:
        q = s.query(DailyReport)
        if target_date:
            q = q.filter(DailyReport.report_date == target_date)
        report = q.order_by(DailyReport.report_date.desc()).first()

    if not report:
        return Response(
            content=b"No report found",
            media_type="text/plain",
            status_code=404,
        )

    sheets = {
        "Overview": [{
            "report_date": str(report.report_date),
            "report_type": report.report_type,
            "n_signals": report.n_signals,
            "n_news": report.n_news,
            "n_top_categories": report.n_top_categories,
            "feishu_sent": report.feishu_sent,
            "created_at": str(report.created_at),
        }],
        "Report": [{"content": report.markdown}],
    }

    filename = f"report_{report.report_date}_{report.report_type}"
    if format == "excel":
        return export_to_excel(sheets, filename)
    elif format == "pdf":
        return export_to_pdf(
            f"Stock Report {report.report_date}",
            [(k, v) for k, v in sheets.items()],
            filename,
        )
    elif format == "docx":
        return export_to_docx(
            f"Stock Report {report.report_date}",
            [(k, v) for k, v in sheets.items()],
            filename,
        )
    else:
        return _export_to_markdown(sheets, filename)


def export_events_list(
    format: str = "excel",
    days_ahead: int = 30,
    industries: str = "",
) -> Response:
    """Export events list to specified format."""
    from datetime import timedelta
    db = get_db()
    today = date.today()
    end = today + timedelta(days=days_ahead)

    industry_list = [i.strip() for i in industries.split(",") if i.strip()]

    with db.session() as s:
        q = s.query(IndustryEvent).filter(
            IndustryEvent.is_future == True,
            IndustryEvent.event_date >= today,
            IndustryEvent.event_date <= end,
        )
        if industry_list:
            q = q.filter(IndustryEvent.industry.in_(industry_list))
        events = q.order_by(IndustryEvent.event_date.asc()).all()

    sheets = {
        "Events": [
            {
                "date": str(e.event_date),
                "title": e.title,
                "industry": e.industry_label,
                "event_type": e.event_type,
                "impact": e.impact_level,
                "related_stocks": e.related_stocks,
                "source": e.source,
            }
            for e in events
        ]
    }
    filename = f"events_{today}_{days_ahead}d"
    if format == "excel":
        return export_to_excel(sheets, filename)
    elif format == "pdf":
        return export_to_pdf(
            f"Events {today} → {end}",
            [("Events", sheets["Events"])],
            filename,
        )
    elif format == "docx":
        return export_to_docx(
            f"Events {today} → {end}",
            [("Events", sheets["Events"])],
            filename,
        )
    else:
        return _export_to_markdown(sheets, filename)


# ============================================================================
# Word (.docx) export
# ============================================================================

def export_to_docx(title: str, sections: list[tuple[str, list[dict]]], filename: str) -> Response:
    """Export to Word using python-docx."""
    if not _has_docx():
        return _export_to_markdown(dict((h, r) for h, r in sections), filename)

    from docx import Document
    from docx.shared import Inches, RGBColor

    doc = Document()
    doc.add_heading(title, level=0)

    for heading, rows in sections:
        doc.add_heading(heading, level=1)
        if not rows:
            doc.add_paragraph("(no data)")
            continue
        headers = list(rows[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        # Header row
        hdr = table.rows[0]
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = str(h)
            for run in cell.paragraphs[0].runs:
                run.bold = True
        # Data rows
        for r in rows[:200]:
            row = table.add_row()
            for i, h in enumerate(headers):
                row.cells[i].text = str(r.get(h, ""))[:200]

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'},
    )